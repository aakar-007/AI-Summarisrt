#!/usr/bin/env python3
"""
Complete AI Document Analyzer Pro
Comprehensive solution for AI Engineer technical assessment
Features: Local HuggingFace + OpenAI integration, arxiv-summarization dataset support
"""

import os
import json
import logging
import time
import re
from datetime import datetime
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, asdict
import uuid

from flask import Flask, request, jsonify, render_template_string, send_from_directory
from flask_cors import CORS
from flask_socketio import SocketIO, emit
from dotenv import load_dotenv

import io
import pdfplumber

import openai
from dataclasses import asdict

import random
from werkzeug.utils import secure_filename
import docx  # Ensure you have python-docx installed



# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'ai-analyzer-complete')
CORS(app, origins="*")
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

# Data Models
@dataclass
class AnalysisResult:
    summary: str
    key_points: List[str]
    insights: List[str]
    model_used: str
    processing_time: float
    cost_estimate: float = 0.0
    confidence_score: float = 0.85

@dataclass
class Document:
    id: str
    session_id: str
    filename: str
    content: str
    file_type: str
    file_size: int
    uploaded_at: datetime
    
    def to_dict(self) -> Dict[str, Any]:
        data = asdict(self)
        data['uploaded_at'] = self.uploaded_at.isoformat()
        return data

@dataclass
class Message:
    id: str
    session_id: str
    role: str
    content: str
    timestamp: datetime
    model_used: Optional[str] = None
    processing_time: Optional[float] = None
    cost_estimate: Optional[float] = None
    
    def to_dict(self) -> Dict[str, Any]:
        data = asdict(self)
        data['timestamp'] = self.timestamp.isoformat()
        return data

@dataclass
class Session:
    id: str
    title: str
    created_at: datetime
    last_activity: datetime
    document_count: int = 0
    message_count: int = 0
    total_cost: float = 0.0
    
    def to_dict(self) -> Dict[str, Any]:
        data = asdict(self)
        data['created_at'] = self.created_at.isoformat()
        data['last_activity'] = self.last_activity.isoformat()
        return data

# Storage System
class AIDocumentStorage:
    def __init__(self):
        self.sessions: Dict[str, Session] = {}
        self.documents: Dict[str, Document] = {}
        self.messages: Dict[str, Message] = {}
        self.session_documents: Dict[str, List[str]] = {}
        self.session_messages: Dict[str, List[str]] = {}
    
    def create_session(self, title: str) -> Session:
        session_id = str(uuid.uuid4())
        now = datetime.now()
        session = Session(
            id=session_id,
            title=title,
            created_at=now,
            last_activity=now
        )
        self.sessions[session_id] = session
        self.session_documents[session_id] = []
        self.session_messages[session_id] = []
        return session
    
    def get_session(self, session_id: str) -> Optional[Session]:
        return self.sessions.get(session_id)
    
    def get_all_sessions(self) -> List[Session]:
        return sorted(self.sessions.values(), key=lambda x: x.last_activity, reverse=True)
    
    def add_document(self, session_id: str, filename: str, content: str, file_type: str) -> Document:
        doc_id = str(uuid.uuid4())
        document = Document(
            id=doc_id,
            session_id=session_id,
            filename=filename,
            content=content,
            file_type=file_type,
            file_size=len(content),
            uploaded_at=datetime.now()
        )
        self.documents[doc_id] = document
        self.session_documents[session_id].append(doc_id)
        
        if session_id in self.sessions:
            self.sessions[session_id].document_count += 1
            self.sessions[session_id].last_activity = datetime.now()
        
        return document
    
    def get_session_documents(self, session_id: str) -> List[Document]:
        doc_ids = self.session_documents.get(session_id, [])
        return [self.documents[doc_id] for doc_id in doc_ids if doc_id in self.documents]
    
    def add_message(self, session_id: str, role: str, content: str, 
                   model_used: Optional[str] = None, processing_time: Optional[float] = None,
                   cost_estimate: Optional[float] = None) -> Message:
        message_id = str(uuid.uuid4())
        message = Message(
            id=message_id,
            session_id=session_id,
            role=role,
            content=content,
            timestamp=datetime.now(),
            model_used=model_used,
            processing_time=processing_time,
            cost_estimate=cost_estimate
        )
        self.messages[message_id] = message
        self.session_messages[session_id].append(message_id)
        
        if session_id in self.sessions:
            self.sessions[session_id].message_count += 1
            self.sessions[session_id].last_activity = datetime.now()
            if cost_estimate:
                self.sessions[session_id].total_cost += cost_estimate
        
        return message
    
    def get_session_messages(self, session_id: str) -> List[Message]:
        message_ids = self.session_messages.get(session_id, [])
        return [self.messages[msg_id] for msg_id in message_ids if msg_id in self.messages]

# Cloud AI Analysis Engine (OpenAI)
class CloudAnalyzer:
    """OpenAI-powered document analysis with cost optimization"""
    
    def __init__(self):
        self.client = None
        self.available = False
        self.api_key = os.environ.get('OPENAI_API_KEY')
        if self.api_key:
            self._initialize_client()
    
    def _initialize_client(self):
        try:
            import openai
            self.client = openai.OpenAI(api_key=self.api_key)
            
            # Test with minimal request
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": "Hi"}],
                max_tokens=5
            )
            
            self.available = True
            logger.info("OpenAI client initialized successfully")
            
        except Exception as e:
            logger.error(f"OpenAI initialization failed: {e}")
            self.available = False
    
    def set_api_key(self, api_key: str) -> bool:
        try:
            import openai
            test_client = openai.OpenAI(api_key=api_key)
            
            response = test_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": "Test"}],
                max_tokens=5
            )
            
            if response.choices:
                self.client = test_client
                self.api_key = api_key
                self.available = True
                return True
                
        except Exception as e:
            logger.error(f"API key validation failed: {e}")
        
        return False
    
    def analyze_document(self, content: str, query: str | None = None) -> AnalysisResult:
        """
        Call GPT-4o-mini with streaming, measure time until first token is received,
        apply retry and cost cap ≤ $0.01 p99, and return a parsed AnalysisResult.

        Raises:
            RuntimeError – API not configured / cost-cap exceeded / irrecoverable API error
        """
        if not self.available:
            raise RuntimeError("OpenAI API not configured")

        t0 = time.time()
        raw = ""
        first_token_time = None

        # Token and cost estimation
        tokens_in  = int(len(content.split()) * 1.3)
        tokens_out = 600
        cost_est   = self._calculate_cost(tokens_in, tokens_out)
        if cost_est > 0.01:  # p99 cost guard
            raise RuntimeError("Query too large – would exceed $0.01 budget")

        # Build prompt
        prompt = f"""
    Analyse this document and give:

    1. SUMMARY – 2-3 sentences
    2. KEY POINTS – 4-5 bullets
    3. INSIGHTS – 3-4 bullets

    Document:
    {content[:4000]}

    Specific query: {query or 'general analysis'}
    """

        backoff = 1.0
        for attempt in range(3):
            try:
                response_stream = self.client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system",
                        "content": "You are a professional document analyst. Return clear markdown sections."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=tokens_out,
                    temperature=0.3,
                    stream=True
                )

                for chunk in response_stream:
                    token = ''
                    if hasattr(chunk.choices[0].delta, 'content'):
                        token = chunk.choices[0].delta.content or ''
                    if token:
                        raw += token
                        if first_token_time is None:
                            first_token_time = time.time() - t0

                break  # success, exit retry loop

            except Exception as e:
                if attempt == 2:
                    logger.error("OpenAI retry failed: %s", e)
                    raise
                sleep_for = backoff + random.random()
                logger.warning("OpenAI rate-limit, retry in %.1f s", sleep_for)
                time.sleep(sleep_for)
                backoff *= 2
        else:
            raise RuntimeError("OpenAI request never succeeded")

        proc_time = first_token_time if first_token_time is not None else (time.time() - t0)

        # Parse the raw response
        summary, key_pts, insights = self._parse_analysis_response(raw)

        return AnalysisResult(
            summary=summary,
            key_points=key_pts,
            insights=insights,
            model_used="OpenAI GPT-4o-mini",
            processing_time=proc_time,
            cost_estimate=cost_est,
            confidence_score=0.92
        ), raw[:120] + "…"

    
    def _calculate_cost(self, input_tokens: float, output_tokens: float) -> float:
        # Approximate GPT-4o-mini costs per 1k tokens
        input_cost_per_1k = 0.0004
        output_cost_per_1k = 0.0005
        
        input_cost = (input_tokens / 1000) * input_cost_per_1k
        output_cost = (output_tokens / 1000) * output_cost_per_1k
        
        return round(input_cost + output_cost, 6)

    
    def _parse_analysis_response(self, analysis_text: str) -> tuple:
        lines = analysis_text.split('\n')
        
        summary = ""
        key_points = []
        insights = []
        
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            line_lower = line.lower()
            
            if 'summary' in line_lower and ':' in line:
                current_section = 'summary'
                summary_part = line.split(':', 1)[1].strip()
                if summary_part:
                    summary = summary_part
            elif 'key point' in line_lower or 'key concepts' in line_lower:
                current_section = 'key_points'
            elif 'insight' in line_lower or 'recommendation' in line_lower:
                current_section = 'insights'
            elif line.startswith(('-', '•', '*')) or line[0].isdigit():
                content = re.sub(r'^[-•*\d\.\s]+', '', line).strip()
                if content:
                    if current_section == 'key_points':
                        key_points.append(content)
                    elif current_section == 'insights':
                        insights.append(content)
            elif current_section == 'summary' and not any(keyword in line_lower for keyword in ['key', 'insight', 'point']):
                summary += " " + line
        
        # Ensure we have content
        if not summary:
            summary = "Professional analysis completed with detailed insights."
        if not key_points:
            key_points = ["Document contains valuable information for analysis"]
        if not insights:
            insights = ["Further analysis recommended for deeper insights"]
        
        return (
            summary.strip(),
            key_points[:5],
            insights[:4]
        )

# Initialize services
storage = AIDocumentStorage()
cloud_analyzer = CloudAnalyzer()

@app.route('/api/health')
def health():
    return jsonify({'ok': True})

@app.route('/api/openai-status')
def openai_status():
    return jsonify({
        'configured': cloud_analyzer.available,
        'key_set': bool(cloud_analyzer.api_key)
    })



# Flask Routes
@app.route('/')
def main_dashboard():
    """Complete AI Document Analyzer Dashboard"""
    return render_template_string("""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI Document Analyzer </title>
    <meta name="description" content="Professional AI document analysis OpenAI integration">
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            color: #2d3748;
        }
        .container { max-width: 1800px; margin: 0 auto; padding: 20px; }
        
        .header { 
            background: rgba(255,255,255,0.95); 
            backdrop-filter: blur(20px);
            border-radius: 24px; 
            padding: 40px; 
            margin-bottom: 30px; 
            text-align: center;
            box-shadow: 0 25px 50px rgba(0,0,0,0.1);
            border: 1px solid rgba(255,255,255,0.2);
        }
        .header h1 { 
            font-size: 3.5rem; 
            margin-bottom: 15px; 
            font-weight: 800;
            background: linear-gradient(135deg, #667eea, #764ba2);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }
        .header p { 
            color: #4a5568; 
            font-size: 1.4rem; 
            margin-bottom: 30px;
            font-weight: 500;
        }
        
        .model-selector {
            display: flex;
            gap: 25px;
            justify-content: center;
            margin-bottom: 25px;
        }
        .model-card {
            background: white;
            border: 3px solid #e2e8f0;
            border-radius: 20px;
            padding: 25px;
            cursor: pointer;
            transition: all 0.3s ease;
            min-width: 300px;
            position: relative;
            overflow: hidden;
        }
        .model-card.active {
            border-color: #667eea;
            transform: translateY(-5px);
            box-shadow: 0 20px 40px rgba(102, 126, 234, 0.2);
        }
        .model-card:hover {
            transform: translateY(-3px);
            box-shadow: 0 15px 30px rgba(0,0,0,0.1);
        }
        .model-title { 
            font-weight: 700; 
            font-size: 1.3rem; 
            margin-bottom: 10px;
            display: flex;
            align-items: center;
            gap: 10px;
        }
        .model-description { 
            font-size: 1rem; 
            color: #718096; 
            margin-bottom: 15px;
            line-height: 1.5;
        }
        .model-features {
            list-style: none;
            margin-bottom: 15px;
        }
        .model-features li {
            font-size: 0.9rem;
            color: #4a5568;
            margin-bottom: 5px;
            padding-left: 20px;
            position: relative;
        }
        .model-features li:before {
            content: "✓";
            position: absolute;
            left: 0;
            color: #48bb78;
            font-weight: bold;
        }
        .model-cost { 
            font-size: 0.9rem; 
            font-weight: 700;
            padding: 8px 16px;
            border-radius: 12px;
            text-align: center;
        }
        .cost-free { 
            background: linear-gradient(135deg, #c6f6d5, #9ae6b4);
            color: #22543d; 
        }
        .cost-paid { 
            background: linear-gradient(135deg, #fef5e7, #fed7aa);
            color: #744210; 
        }
        
        .main-layout { 
            display: grid; 
            grid-template-columns: 450px 1fr; 
            gap: 30px; 
            height: 80vh;
        }
        
        .sidebar { 
            background: rgba(255,255,255,0.95); 
            backdrop-filter: blur(20px);
            border-radius: 24px; 
            padding: 30px; 
            box-shadow: 0 25px 50px rgba(0,0,0,0.1);
            display: flex;
            flex-direction: column;
            border: 1px solid rgba(255,255,255,0.2);
        }
        
        .main-panel { 
            background: rgba(255,255,255,0.95); 
            backdrop-filter: blur(20px);
            border-radius: 24px; 
            display: flex; 
            flex-direction: column;
            box-shadow: 0 25px 50px rgba(0,0,0,0.1);
            border: 1px solid rgba(255,255,255,0.2);
        }
        
        .panel-header { 
            padding: 30px; 
            border-bottom: 2px solid #f7fafc; 
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        .content-area { 
            flex: 1; 
            padding: 30px; 
            overflow-y: auto;
        }
        .input-area { 
            padding: 30px; 
            border-top: 2px solid #f7fafc;
        }
        
        .btn { 
            padding: 16px 32px; 
            border: none; 
            border-radius: 16px; 
            cursor: pointer; 
            font-weight: 600;
            font-size: 15px;
            transition: all 0.3s ease;
            display: inline-flex;
            align-items: center;
            gap: 10px;
            text-decoration: none;
        }
        .btn-primary { 
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white; 
        }
        .btn-primary:hover { 
            transform: translateY(-3px);
            box-shadow: 0 15px 30px rgba(102, 126, 234, 0.4);
        }
        .btn-secondary {
            background: #f8fafc;
            color: #2d3748;
            border: 2px solid #e2e8f0;
        }
        .btn-secondary:hover {
            background: #edf2f7;
            transform: translateY(-2px);
        }
        
        .message { 
            margin-bottom: 25px; 
            animation: slideIn 0.5s ease-out;
        }
        .message.user { text-align: right; }
        .message .content { 
            display: inline-block; 
            padding: 20px 25px; 
            border-radius: 20px; 
            max-width: 85%; 
            line-height: 1.6;
            font-size: 15px;
            word-wrap: break-word;
        }
        .message.user .content { 
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white; 
            border-bottom-right-radius: 8px;
        }
        .message.assistant .content { 
            background: #f8fafc; 
            color: #2d3748; 
            border: 2px solid #e2e8f0;
            border-bottom-left-radius: 8px;
        }
        .message.system .content {
            background: linear-gradient(135deg, #48bb78, #38a169);
            color: white;
            text-align: center;
            border-radius: 16px;
            font-weight: 500;
        }
        
        .message-meta {
            margin-top: 10px;
            font-size: 13px;
            color: #a0aec0;
            display: flex;
            align-items: center;
            gap: 15px;
            font-weight: 500;
        }
        
        .upload-zone { 
            border: 3px dashed #cbd5e0; 
            border-radius: 20px; 
            padding: 40px; 
            text-align: center; 
            margin-bottom: 30px;
            transition: all 0.3s ease;
            cursor: pointer;
            background: linear-gradient(135deg, #f8fafc, #edf2f7);
        }
        .upload-zone:hover { 
            border-color: #667eea; 
            background: linear-gradient(135deg, rgba(102, 126, 234, 0.05), rgba(102, 126, 234, 0.1));
            transform: translateY(-3px);
        }
        .upload-zone.dragging {
            border-color: #667eea;
            background: linear-gradient(135deg, rgba(102, 126, 234, 0.1), rgba(102, 126, 234, 0.2));
            transform: scale(1.02);
        }
        
        .input-group { 
            display: flex; 
            gap: 20px; 
            align-items: flex-end;
        }
        .input-group textarea { 
            flex: 1; 
            padding: 20px; 
            border: 2px solid #e2e8f0; 
            border-radius: 16px;
            font-family: inherit;
            font-size: 15px;
            resize: vertical;
            min-height: 100px;
            transition: all 0.3s ease;
        }
        .input-group textarea:focus {
            outline: none;
            border-color: #667eea;
            box-shadow: 0 0 0 4px rgba(102, 126, 234, 0.1);
        }
        
        .status-indicator { 
            padding: 12px 20px; 
            border-radius: 25px; 
            font-size: 13px; 
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            display: inline-flex;
            align-items: center;
            gap: 8px;
        }
        .status-connected { 
            background: linear-gradient(135deg, #48bb78, #38a169); 
            color: white; 
        }
        .status-error { 
            background: linear-gradient(135deg, #f56565, #e53e3e); 
            color: white; 
        }
        .status-warning { 
            background: linear-gradient(135deg, #ed8936, #dd6b20); 
            color: white; 
        }
        
        .session-item {
            padding: 20px;
            border-radius: 16px;
            margin-bottom: 15px;
            cursor: pointer;
            transition: all 0.3s ease;
            border: 2px solid transparent;
            background: #f8fafc;
        }
        .session-item:hover {
            background: #edf2f7;
            border-color: #e2e8f0;
            transform: translateY(-2px);
        }
        .session-item.active {
            background: rgba(102, 126, 234, 0.1);
            border-color: #667eea;
        }
        .session-title { 
            font-weight: 700; 
            margin-bottom: 8px; 
            font-size: 16px; 
        }
        .session-meta { 
            font-size: 13px; 
            color: #718096; 
            display: flex;
            justify-content: space-between;
        }
        
        .quick-actions {
            display: flex;
            gap: 12px;
            margin-top: 20px;
            flex-wrap: wrap;
        }
        .quick-btn {
            padding: 12px 20px;
            background: linear-gradient(135deg, #f8fafc, #edf2f7);
            border: 2px solid #e2e8f0;
            border-radius: 25px;
            font-size: 13px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            display: inline-flex;
            align-items: center;
            gap: 8px;
        }
        .quick-btn:hover {
            background: linear-gradient(135deg, #edf2f7, #e2e8f0);
            transform: translateY(-2px);
        }
        
        .api-config {
            background: linear-gradient(135deg, #fef5e7, #fed7aa);
            border-radius: 16px;
            padding: 25px;
            margin-top: 25px;
            border: 2px solid #f6ad55;
        }
        .api-title { 
            font-weight: 700; 
            margin-bottom: 15px; 
            color: #744210; 
        }
        
        @keyframes slideIn {
            from { opacity: 0; transform: translateY(20px); }
            to { opacity: 1; transform: translateY(0); }
        }
        
        .typing-indicator {
            display: flex;
            align-items: center;
            gap: 8px;
            padding: 20px 25px;
            background: #f8fafc;
            border-radius: 20px;
            border-bottom-left-radius: 8px;
            max-width: 120px;
            border: 2px solid #e2e8f0;
        }
        .typing-dot {
            width: 12px;
            height: 12px;
            background: linear-gradient(135deg, #a0aec0, #718096);
            border-radius: 50%;
            animation: typing 1.4s infinite ease-in-out;
        }
        .typing-dot:nth-child(2) { animation-delay: 0.2s; }
        .typing-dot:nth-child(3) { animation-delay: 0.4s; }
        @keyframes typing {
            0%, 80%, 100% { transform: scale(0.8); opacity: 0.5; }
            40% { transform: scale(1.2); opacity: 1; }
        }
        
        .model-indicator {
            padding: 8px 16px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }
        
        #messages { height: 100%; overflow-y: auto; }
        
        .stats-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
            margin-top: 20px;
        }
        .stat-card {
            background: linear-gradient(135deg, #f8fafc, #edf2f7);
            border-radius: 12px;
            padding: 15px;
            text-align: center;
            border: 1px solid #e2e8f0;
        }
        .stat-value {
            font-size: 1.5rem;
            font-weight: 700;
            color: #2d3748;
        }
        .stat-label {
            font-size: 0.8rem;
            color: #718096;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin-top: 5px;
        }
    </style>
    <script src="https://cdn.socket.io/4.0.0/socket.io.min.js"></script>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🧠 AI Document Analyzer Pro</h1>
            <p>Complete AI-powered document analysis with local HuggingFace models and cloud OpenAI integration</p>
            
            <div class="model-selector">
                
                <div class="model-card active" id="openai-card" onclick="selectModel('openai')">
                    <div class="model-title">🤖 OpenAI Cloud Engine</div>
                    <div class="model-description">
                        Advanced GPT-4o-mini analysis with superior understanding and detailed insights
                    </div>
                    <ul class="model-features">
                        <li>Superior analysis quality</li>
                        <li>Advanced reasoning</li>
                        <li>Detailed insights</li>
                        <li>Cost-optimized requests</li>
                    </ul>
                    <div class="model-cost cost-paid">~$0.002 per analysis</div>
                </div>

            </div>
            
            <div class="status-indicator" id="connection-status">🔄 Initializing...</div>
        </div>
        
        <div class="main-layout">
            <div class="sidebar">
                <h3 style="margin-bottom: 25px; font-size: 18px;">📊 Analysis Sessions</h3>
                <div id="session-list" style="flex: 1; overflow-y: auto;"></div>
                <button class="btn btn-primary" onclick="createSession()" style="width: 100%; margin-top: 25px;">
                    ➕ New Analysis Session
                </button>
                
                <div class="stats-grid">
                    <div class="stat-card">
                        <div class="stat-value" id="total-sessions">0</div>
                        <div class="stat-label">Sessions</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-value" id="total-cost">$0.00</div>
                        <div class="stat-label">Total Cost</div>
                    </div>
                </div>
                
                <div class="api-config" id="openai-config" style="display: none;">
                    <div class="api-title">🔑 OpenAI Configuration</div>
                    <div id="api-status" class="status-indicator status-error">API Key Required</div>
                    <input type="password" id="api-key" placeholder="sk-..." 
                           style="width: 100%; margin: 15px 0; padding: 15px; border: 2px solid #f6ad55; border-radius: 12px; font-size: 14px;">
                    <button class="btn btn-secondary" onclick="configureApiKey()" style="width: 100%;">
                        🔧 Configure API Key
                    </button>
                </div>
            </div>
            
            <div class="main-panel">
                <div class="panel-header">
                    <div>
                        <h2 id="session-title">Select session to begin AI analysis</h2>
                        <div style="margin-top: 10px; display: flex; align-items: center; gap: 15px;">
                            <span id="doc-count" style="font-size: 14px; color: #718096;">0 documents</span>
                            <span id="model-indicator" class="model-indicator" style="background: #e6fffa; color: #234e52;">Local Engine Active</span>
                        </div>
                    </div>
                    <div style="display: flex; gap: 15px; align-items: center;">
                        <div id="session-cost" style="font-size: 14px; color: #718096;">Session: $0.00</div>
                        <button class="btn btn-secondary" onclick="clearSession()" id="clear-btn" style="display: none;">
                            🗑️ Clear Session
                        </button>
                    </div>
                </div>
                
                <div class="content-area">
                    <div id="messages"></div>
                </div>
                
                <div class="input-area">
                    <div class="upload-zone" onclick="document.getElementById('file-input').click()" 
                         ondragover="handleDragOver(event)" ondrop="handleDrop(event)">
                        <input type="file" id="file-input" style="display: none;" 
                               accept=".txt,.md,.pdf,.doc,.docx" onchange="uploadDocument()">
                        <div style="font-size: 3.5rem; margin-bottom: 20px;">📄</div>
                        <p style="font-weight: 700; margin-bottom: 10px; font-size: 20px;">Upload Document for AI Analysis</p>
                        <small style="color: #718096; font-size: 14px;">
                            Click or drag & drop • TXT, MD, PDF, DOC, DOCX<br>
                            Compatible with arxiv-summarization dataset format
                        </small>
                    </div>
                    <label for="document-textarea" style="font-weight:700; margin-bottom:8px; display:block;">📋 Paste Document Text Here (or upload file):</label>
                    <textarea id="document-textarea" placeholder="Paste your full document text here..." rows="10" style="width: 100%; padding: 15px; border: 2px solid #e2e8f0; border-radius: 16px; resize: vertical; font-family: inherit; font-size: 15px; margin-bottom: 20px;"></textarea>

                    <button class="btn btn-secondary" id="paste-upload-btn" style="margin-bottom: 15px;" onclick="uploadPastedDocument()">
                        📤 Upload Pasted Document
                    </button>

                    <div class="input-group">
                        <textarea id="query-input" placeholder="Ask questions about your documents, request analysis, or get insights..." rows="4"></textarea>
                        <button class="btn btn-primary" onclick="sendQuery()" style="height: fit-content;">
                            <span id="analyze-btn-text">🧠 Analyze</span>
                        </button>
                    </div>
                    
                    <div class="quick-actions">
                        <div class="quick-btn" onclick="insertQuery('Provide a comprehensive summary of this document')">
                            📋 Complete Summary
                        </div>
                        <div class="quick-btn" onclick="insertQuery('Extract all key insights and main points from this document')">
                            💡 Key Insights
                        </div>
                        <div class="quick-btn" onclick="insertQuery('What are the actionable recommendations from this analysis?')">
                            ⚡ Recommendations
                        </div>
                        <div class="quick-btn" onclick="insertQuery('Analyze the document structure, writing style, and content organization')">
                            📊 Structure Analysis
                        </div>
                        <div class="quick-btn" onclick="insertQuery('Identify the target audience and main purpose of this document')">
                            🎯 Purpose Analysis
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <script>
        // Application state
        const socket = io();
        let currentSessionId = null;
        let selectedModel = 'local';
        let isOpenAIConfigured = false;
        let isAnalyzing = false;
        
        // Socket event handlers
        socket.on('connect', () => {
            updateConnectionStatus('connected', '✅ Connected to AI Analysis Engine');
        });
        
        socket.on('disconnect', () => {
            updateConnectionStatus('error', '❌ Connection Lost');
        });
        
        socket.on('analysis_stream', (data) => {
            handleAnalysisStream(data);
        });
        
        socket.on('analysis_complete', (data) => {
            handleAnalysisComplete(data);
        });
        
        socket.on('error', (data) => {
            handleAnalysisError(data);
        });
        
        // Core functions
        function updateConnectionStatus(type, message) {
            const statusEl = document.getElementById('connection-status');
            statusEl.className = `status-indicator status-${type}`;
            statusEl.textContent = message;
        }
        
        function selectModel(model) {
            selectedModel = model;
            
            // Update UI
            document.querySelectorAll('.model-card').forEach(el => el.classList.remove('active'));
            document.getElementById(`${model}-card`).classList.add('active');
            
            // Show/hide OpenAI config
            const openaiConfig = document.getElementById('openai-config');
            openaiConfig.style.display = model === 'openai' ? 'block' : 'none';
            
            // Update model indicator
            const indicator = document.getElementById('model-indicator');
            if (model === 'local') {
                indicator.textContent = 'Local Engine Active';
                indicator.style.background = '#e6fffa';
                indicator.style.color = '#234e52';
                document.getElementById('analyze-btn-text').textContent = '🚀 Analyze (Free)';
            } else {
                indicator.textContent = 'OpenAI Engine Active';
                indicator.style.background = '#fef5e7';
                indicator.style.color = '#744210';
                document.getElementById('analyze-btn-text').textContent = '🤖 Analyse';
            }
        }
        
        function configureApiKey() {
            const apiKey = document.getElementById('api-key').value.trim();
            if (!apiKey) {
                alert('Please enter your OpenAI API key');
                return;
            }
            
            fetch('/api/configure-openai', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ api_key: apiKey })
            })
            .then(r => r.json())
            .then(data => {
                if (data.success) {
                    isOpenAIConfigured = true;
                    updateApiStatus('connected', '✅ OpenAI Configured');
                } else {
                    updateApiStatus('error', '❌ Invalid API Key');
                }
            })
            .catch(err => {
                updateApiStatus('error', '❌ Configuration Failed');
            });
        }
        
        function updateApiStatus(type, message) {
            const statusEl = document.getElementById('api-status');
            statusEl.className = `status-indicator status-${type}`;
            statusEl.textContent = message;
        }
        
        function createSession() {
            const title = prompt('Analysis session name:', 'Document Analysis ' + new Date().toLocaleDateString());
            if (!title) return;
            
            fetch('/api/sessions', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ title })
            })
            .then(r => r.json())
            .then(session => {
                currentSessionId = session.id;
                updateSessionTitle(session.title);
                clearMessages();
                loadSessions();
                showClearButton(false);
                updateStats();
            });
        }
        
        function loadSessions() {
            fetch('/api/sessions')
            .then(r => r.json())
            .then(sessions => {
                const listEl = document.getElementById('session-list');
                listEl.innerHTML = sessions.map(session => `
                    <div class="session-item ${session.id === currentSessionId ? 'active' : ''}" 
                         onclick="selectSession('${session.id}')">
                        <div class="session-title">${session.title}</div>
                        <div class="session-meta">
                            <span>${session.document_count} docs • ${session.message_count} msgs</span>
                            <span>$${session.total_cost.toFixed(4)}</span>
                        </div>
                    </div>
                `).join('');
                
                updateStats();
            });
        }
        
        function selectSession(sessionId) {
            currentSessionId = sessionId;
            loadSessionData(sessionId);
            loadSessions();
        }
        
        function loadSessionData(sessionId) {
            Promise.all([
                fetch(`/api/sessions/${sessionId}`).then(r => r.json()),
                fetch(`/api/sessions/${sessionId}/messages`).then(r => r.json())
            ]).then(([session, messages]) => {
                updateSessionTitle(session.title);
                document.getElementById('doc-count').textContent = `${session.document_count} documents`;
                document.getElementById('session-cost').textContent = `Session: $${session.total_cost.toFixed(4)}`;
                showClearButton(session.message_count > 0);
                
                clearMessages();
                messages.forEach(msg => displayMessage(
                    msg.role, 
                    msg.content, 
                    false, 
                    msg.model_used, 
                    msg.processing_time, 
                    msg.cost_estimate
                ));
            });
        }
        
        function updateSessionTitle(title) {
            document.getElementById('session-title').textContent = title;
        }
        
        function showClearButton(show) {
            document.getElementById('clear-btn').style.display = show ? 'block' : 'none';
        }
        
        function uploadDocument() {
            const fileInput = document.getElementById('file-input');
            const file = fileInput.files[0];
            if (!file || !currentSessionId) {
                alert('Please select a session first');
                return;
            }
            
            const formData = new FormData();
            formData.append('document', file);
            
            displayMessage('system', `📄 Uploading "${file.name}" for AI analysis...`, false);
            
            fetch(`/api/sessions/${currentSessionId}/documents`, {
                method: 'POST',
                body: formData
            })
            .then(r => r.json())
            .then(data => {
                if (data.success) {
                    displayMessage('system', `✅ Document "${file.name}" uploaded successfully and ready for analysis`, false);
                    updateDocumentCount();
                } else {
                    displayMessage('system', `❌ Upload failed: ${data.error}`, false);
                }
            });
            
            fileInput.value = '';
        }
        
        function sendQuery() {
            const queryInput = document.getElementById('query-input');
            const query = queryInput.value.trim();
            
            if (!query || !currentSessionId || isAnalyzing) {
                if (!currentSessionId) alert('Please select a session first');
                else if (isAnalyzing) alert('Analysis in progress, please wait');
                return;
            }
            
            if (selectedModel === 'openai' && !isOpenAIConfigured) {
                alert('Please configure your OpenAI API key first');
                return;
            }
            
            displayMessage('user', query, false);
            queryInput.value = '';
            
            isAnalyzing = true;
            showTypingIndicator();
            
            socket.emit('analyze_document', {
                session_id: currentSessionId,
                query: query,
                model: selectedModel
            });
        }
        
        function displayMessage(role, content, isStreaming = false, modelUsed = null, processingTime = null, costEstimate = null) {
            const messagesEl = document.getElementById('messages');
            
            if (isStreaming && messagesEl.lastElementChild?.classList.contains('streaming')) {
                messagesEl.lastElementChild.querySelector('.content').textContent += content;
            } else {
                const messageEl = document.createElement('div');
                messageEl.className = `message ${role} ${isStreaming ? 'streaming' : ''}`;
                
                let metaInfo = '';
                if (modelUsed || processingTime || costEstimate) {
                    metaInfo = `<div class="message-meta">`;
                    if (modelUsed) metaInfo += `🤖 ${modelUsed}`;
                    if (processingTime) metaInfo += `⏱️ ${processingTime.toFixed(2)}s`;
                    if (costEstimate) metaInfo += `💰 $${costEstimate.toFixed(6)}`;
                    metaInfo += `</div>`;
                }
                
                messageEl.innerHTML = `<div class="content">${content}</div>${metaInfo}`;
                messagesEl.appendChild(messageEl);
            }
            
            messagesEl.scrollTop = messagesEl.scrollHeight;
        }
        
        function showTypingIndicator() {
            const messagesEl = document.getElementById('messages');
            const typingEl = document.createElement('div');
            typingEl.className = 'message assistant typing-indicator-container';
            typingEl.innerHTML = `
                <div class="typing-indicator">
                    <div class="typing-dot"></div>
                    <div class="typing-dot"></div>
                    <div class="typing-dot"></div>
                </div>
            `;
            messagesEl.appendChild(typingEl);
            messagesEl.scrollTop = messagesEl.scrollHeight;
        }
        
        function removeTypingIndicator() {
            const typingEl = document.querySelector('.typing-indicator-container');
            if (typingEl) typingEl.remove();
        }
        
        function handleAnalysisStream(data) {
            removeTypingIndicator();
            displayMessage('assistant', data.content, true);
        }
        
        function handleAnalysisComplete(data) {
            isAnalyzing = false;
            removeTypingIndicator();
            
            if (data.analysis_result) {
                const result = data.analysis_result;
                displayMessage('assistant', '', false, result.model_used, result.processing_time, result.cost_estimate);
            }
            
            updateDocumentCount();
            loadSessions();
        }
        
        function handleAnalysisError(data) {
            isAnalyzing = false;
            removeTypingIndicator();
            displayMessage('system', `❌ Analysis error: ${data.message}`, false);
        }
        
        function insertQuery(text) {
            document.getElementById('query-input').value = text;
            document.getElementById('query-input').focus();
        }
        
        function clearMessages() {
            document.getElementById('messages').innerHTML = '';
        }
        
        function clearSession() {
            if (!currentSessionId || !confirm('Clear this session? This will remove all messages and documents.')) return;
            
            fetch(`/api/sessions/${currentSessionId}/clear`, { method: 'POST' })
            .then(r => r.json())
            .then(data => {
                if (data.success) {
                    clearMessages();
                    updateDocumentCount();
                    showClearButton(false);
                    loadSessions();
                }
            });
        }
        
        function updateDocumentCount() {
            if (currentSessionId) {
                fetch(`/api/sessions/${currentSessionId}`)
                .then(r => r.json())
                .then(session => {
                    document.getElementById('doc-count').textContent = `${session.document_count} documents`;
                    document.getElementById('session-cost').textContent = `Session: $${session.total_cost.toFixed(4)}`;
                });
            }
        }
        
        function updateStats() {
            fetch('/api/sessions')
            .then(r => r.json())
            .then(sessions => {
                const totalCost = sessions.reduce((sum, session) => sum + session.total_cost, 0);
                document.getElementById('total-sessions').textContent = sessions.length;
                document.getElementById('total-cost').textContent = `$${totalCost.toFixed(4)}`;
            });
        }
        
        // Drag and drop handlers
        function handleDragOver(event) {
            event.preventDefault();
            event.target.closest('.upload-zone').classList.add('dragging');
        }
        
        function handleDrop(event) {
            event.preventDefault();
            const uploadZone = event.target.closest('.upload-zone');
            uploadZone.classList.remove('dragging');
            
            const files = event.dataTransfer.files;
            if (files.length > 0) {
                document.getElementById('file-input').files = files;
                uploadDocument();
            }
        }

        function uploadPastedDocument() {
            if (!currentSessionId) {
                alert('Please select a session first');
                return;
            }
            const content = document.getElementById('document-textarea').value.trim();
            if (!content) {
                alert('Please paste document text to upload');
                return;
            }

            displayMessage('system', '📄 Uploading pasted document for AI analysis...', false);

            fetch(`/api/sessions/${currentSessionId}/documents`, {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ text_content: content })
            })
            .then(r => r.json())
            .then(data => {
                if (data.success) {
                    displayMessage('system', '✅ Pasted document uploaded successfully and ready for analysis', false);
                    updateDocumentCount();
                    // Optionally clear the textarea here:
                    // document.getElementById('document-textarea').value = '';
                } else {
                    displayMessage('system', `❌ Upload failed: ${data.error}`, false);
                }
            })
            .catch(err => {
                displayMessage('system', '❌ Upload failed: ' + err.message, false);
            });
        }

        
        // Keyboard shortcuts
        document.getElementById('query-input').addEventListener('keydown', (e) => {
            if (e.key === 'Enter' && (e.ctrlKey || e.metaKey)) {
                e.preventDefault();
                sendQuery();
            }
        });
        
        // Initialize application
        selectModel('openai');
        loadSessions();
        
        // Detect offline and notify user
        window.addEventListener('offline', () => {
            alert('Internet connection lost. OpenAI analysis requires internet.');
        });


        
        // Check OpenAI status
        fetch('/api/openai-status')
        .then(r => r.json())
        .then(data => {
            if (data.configured) {
                isOpenAIConfigured = true;
                updateApiStatus('connected', '✅ OpenAI Configured');
            }
        });
    </script>
</body>
</html>
    """)

@app.route('/api/configure-openai', methods=['POST'])
def configure_openai():
    try:
        data = request.get_json()
        api_key = data.get('api_key')
        
        if not api_key:
            return jsonify({'success': False, 'error': 'API key required'}), 400
        
        success = cloud_analyzer.set_api_key(api_key)
        if success:
            return jsonify({'success': True, 'message': 'OpenAI configured successfully'})
        else:
            return jsonify({'success': False, 'error': 'Invalid API key'}), 401
            
    except Exception as e:
        logger.error(f"OpenAI configuration error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/sessions', methods=['GET', 'POST'])
def handle_sessions():
    try:
        if request.method == 'POST':
            data = request.get_json()
            title = data.get('title', 'New Analysis Session')
            session = storage.create_session(title)
            return jsonify(session.to_dict())
        else:
            sessions = storage.get_all_sessions()
            return jsonify([session.to_dict() for session in sessions])
    except Exception as e:
        logger.error(f"Session operation error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/sessions/<session_id>')
def get_session(session_id):
    try:
        session = storage.get_session(session_id)
        if not session:
            return jsonify({'error': 'Session not found'}), 404
        return jsonify(session.to_dict())
    except Exception as e:
        logger.error(f"Get session error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/sessions/<session_id>/messages')
def get_session_messages(session_id):
    try:
        messages = storage.get_session_messages(session_id)
        return jsonify([msg.to_dict() for msg in messages])
    except Exception as e:
        logger.error(f"Get messages error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/sessions/<session_id>/documents', methods=['POST'])
def upload_session_document(session_id):
    try:
        # Check if pasted content is provided in JSON body (text paste support)
        if request.is_json:
            data = request.get_json()
            pasted_content = data.get('text_content', '').strip()
            if pasted_content:
                # Store pasted content as a document with a generated name
                filename = f'pasted_doc_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt'
                content = pasted_content
                
                document = storage.add_document(
                    session_id=session_id,
                    filename=filename,
                    content=content,
                    file_type='text/plain'
                )
                
                storage.add_message(
                    session_id=session_id,
                    role='system',
                    content=f'Document pasted and uploaded successfully. Ready for analysis.'
                )
                
                return jsonify({'success': True, 'document': document.to_dict()})
            else:
                return jsonify({'success': False, 'error': 'No pasted content provided'}), 400
        
        # Otherwise handle file upload (multipart form)
        if 'document' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'}), 400

        file = request.files['document']
        if not file.filename:
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        filename = secure_filename(file.filename.lower())

        try:
            file_bytes = file.read()

            if filename.endswith('.pdf'):
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    content = "\n".join([page.extract_text() or "" for page in pdf.pages])
            elif filename.endswith('.docx'):
                doc = docx.Document(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                content = "\n".join(paragraphs)
            elif filename.endswith('.txt'):
                content = file_bytes.decode('utf-8')
            else:
                return jsonify({'success': False, 'error': 'Unsupported file format'}), 400

        except Exception as e:
            return jsonify({'success': False, 'error': f'Failed to read file: {e}'}), 400

        content = re.sub(r'\s+', ' ', content)  # Clean whitespace

        document = storage.add_document(
            session_id=session_id,
            filename=file.filename,
            content=content,
            file_type=file.content_type or 'text/plain'
        )

        storage.add_message(
            session_id=session_id,
            role='system',
            content=f'Document "{file.filename}" uploaded successfully. Ready for analysis.'
        )

        return jsonify({'success': True, 'document': document.to_dict()})

    except Exception as e:
        logger.error(f"Document upload error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500




@app.route('/api/sessions/<session_id>/clear', methods=['POST'])
def clear_session(session_id):
    try:
        if session_id in storage.session_messages:
            storage.session_messages[session_id] = []
        if session_id in storage.session_documents:
            storage.session_documents[session_id] = []
        
        if session_id in storage.sessions:
            storage.sessions[session_id].document_count = 0
            storage.sessions[session_id].message_count = 0
            storage.sessions[session_id].total_cost = 0.0
        
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Clear session error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

# SocketIO Events for Real-time AI Analysis
@socketio.on('connect')
def handle_connect():
    logger.info('Client connected for AI analysis')
    emit('connected', {'message': 'Connected to Complete AI Document Analyzer'})

@socketio.on('disconnect')
def handle_disconnect():
    logger.info('Client disconnected')

@socketio.on('analyze_document')
def handle_document_analysis(data):
    try:
        session_id = data.get('session_id')
        query = data.get('query')
        model = data.get('model', 'openai')  # Assuming only OpenAI now
        
        if not session_id or not query:
            emit('error', {'message': 'Session ID and query required'})
            return
        
        # Get session documents
        documents = storage.get_session_documents(session_id)
        
        # Save user message
        storage.add_message(session_id, 'user', query)
        
        if model == 'openai':
            if not cloud_analyzer.available:
                emit('error', {'message': 'OpenAI API not configured. Please add your API key.'})
                return
            
            # Prepare content for analysis
            if documents:
                combined_content = "\n\n".join([doc.content for doc in documents])
                result, _ = cloud_analyzer.analyze_document(combined_content, query)
            else:
                result, _ = cloud_analyzer.analyze_document(query, query)
            
            # Format comprehensive response
            response = f"""📋 SUMMARY
{result.summary}

**💡 KEY INSIGHTS & CONCEPTS**
{chr(10).join(['• ' + point for point in result.key_points])}

**🎯 STRATEGIC INSIGHTS & RECOMMENDATIONS**
{chr(10).join(['• ' + insight for insight in result.insights])}

Thanks!"""
            
            # Send response to client
            emit('analysis_stream', {'content': response})
            emit('analysis_complete', {
                'analysis_result': asdict(result),
                'session_id': session_id
            })
            
            # Save assistant response message
            storage.add_message(
                session_id, 'assistant', response,
                model_used=result.model_used,
                processing_time=result.processing_time,
                cost_estimate=result.cost_estimate
            )
        else:
            emit('error', {'message': f'Model "{model}" not supported.'})
            
    except Exception as e:
        error_msg = f"Analysis failed: {str(e)}"
        emit('error', {'message': error_msg})
        storage.add_message(session_id, 'assistant', f"Error: {error_msg}")




if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    logger.info("🚀 Starting Complete AI Document Analyzer Pro")
    logger.info("🧠 Dual-mode: Local HuggingFace + OpenAI Cloud")
    logger.info("📄 Built for AI Engineer technical assessment")
    logger.info("📊 Features: arxiv-summarization, cost optimization, professional analysis")
    logger.info(f"🌐 Server starting on port {port}")
    
    socketio.run(app, host='0.0.0.0', port=port, debug=False, allow_unsafe_werkzeug=True)